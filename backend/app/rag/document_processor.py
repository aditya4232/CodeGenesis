"""
MemoRAG ULTRA - Document Handlers
Extract text from various document formats
"""
import io
from pathlib import Path
from typing import Optional, Dict, Any, List
import logging

# PDF processing
try:
    import pypdf
    from pdfplumber import open as pdf_open
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF libraries not available")

# DOCX processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

# HTML processing
try:
    from bs4 import BeautifulSoup
    import requests
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logging.warning("BeautifulSoup not available")

logger = logging.getLogger(__name__)


class DocumentHandler:
    """Base document handler"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from document
        
        Args:
            file_path: Path to document
            
        Returns:
            Dict with 'text' and 'metadata'
        """
        raise NotImplementedError


class PDFHandler(DocumentHandler):
    """PDF document handler"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF libraries not installed")
        
        text_parts = []
        metadata = {
            'pages': 0,
            'has_images': False,
            'has_tables': False
        }
        
        try:
            # Use pdfplumber for better text extraction
            with pdf_open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                    
                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['has_tables'] = True
                        for table in tables:
                            # Convert table to text
                            table_text = '\n'.join([
                                ' | '.join([str(cell) if cell else '' for cell in row])
                                for row in table
                            ])
                            text_parts.append(f"\n[Table on Page {page_num}]\n{table_text}\n")
                    
                    # Check for images
                    if page.images:
                        metadata['has_images'] = True
            
            full_text = '\n\n'.join(text_parts)
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise


class TextHandler(DocumentHandler):
    """Plain text handler"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                'text': text,
                'metadata': {
                    'lines': text.count('\n') + 1,
                    'chars': len(text)
                }
            }
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            
            return {
                'text': text,
                'metadata': {
                    'lines': text.count('\n') + 1,
                    'chars': len(text),
                    'encoding': 'latin-1'
                }
            }


class MarkdownHandler(DocumentHandler):
    """Markdown handler"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Count headers
        headers = text.count('#')
        
        return {
            'text': text,
            'metadata': {
                'headers': headers,
                'lines': text.count('\n') + 1
            }
        }


class DOCXHandler(DocumentHandler):
    """DOCX document handler"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        try:
            doc = docx.Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            table_count = len(doc.tables)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    table_text.append(row_text)
                text_parts.append('\n[Table]\n' + '\n'.join(table_text))
            
            full_text = '\n\n'.join(text_parts)
            
            return {
                'text': full_text,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': table_count
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise


class URLHandler(DocumentHandler):
    """URL/Web page handler"""
    
    def extract_text(self, url: str) -> Dict[str, Any]:
        """Extract text from URL"""
        if not HTML_AVAILABLE:
            raise RuntimeError("BeautifulSoup not installed")
        
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Get title
            title = soup.title.string if soup.title else url
            
            return {
                'text': text,
                'metadata': {
                    'url': url,
                    'title': title,
                    'content_type': response.headers.get('content-type', 'unknown')
                }
            }
            
        except Exception as e:
            logger.error(f"URL extraction failed: {e}")
            raise


class DocumentProcessor:
    """Main document processor that routes to appropriate handler"""
    
    def __init__(self):
        self.handlers = {
            '.pdf': PDFHandler(),
            '.txt': TextHandler(),
            '.md': MarkdownHandler(),
            '.docx': DOCXHandler(),
        }
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process file and extract text
        
        Args:
            file_path: Path to file
            
        Returns:
            Dict with 'text' and 'metadata'
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.handlers:
            raise ValueError(f"Unsupported file type: {ext}")
        
        handler = self.handlers[ext]
        result = handler.extract_text(str(path))
        
        # Add file metadata
        result['metadata']['file_name'] = path.name
        result['metadata']['file_size'] = path.stat().st_size
        result['metadata']['file_type'] = ext
        
        return result
    
    def process_url(self, url: str) -> Dict[str, Any]:
        """
        Process URL and extract text
        
        Args:
            url: URL to process
            
        Returns:
            Dict with 'text' and 'metadata'
        """
        handler = URLHandler()
        return handler.extract_text(url)
    
    def process_content(self, content: str, content_type: str = "text") -> Dict[str, Any]:
        """
        Process raw content
        
        Args:
            content: Raw content string
            content_type: Type of content
            
        Returns:
            Dict with 'text' and 'metadata'
        """
        return {
            'text': content,
            'metadata': {
                'content_type': content_type,
                'chars': len(content)
            }
        }


def get_document_processor() -> DocumentProcessor:
    """Get document processor instance"""
    return DocumentProcessor()
